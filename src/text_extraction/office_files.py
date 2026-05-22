from __future__ import annotations

import subprocess

from src.models import ExtractionResult, FileRecord
from src.utils.text import normalize_text, truncate_text


def extract_office(record: FileRecord, max_rows: int) -> ExtractionResult:
    if record.extension == ".docx":
        return extract_docx(record)
    if record.extension == ".rtf":
        return extract_rtf(record)
    if record.extension in {".xlsx", ".xls"}:
        return extract_excel(record, max_rows=max_rows)
    if record.extension == ".doc":
        return extract_doc(record)
    return ExtractionResult("", "skipped", warnings=[f"Unsupported office extension: {record.extension}"])


def extract_docx(record: FileRecord) -> ExtractionResult:
    try:
        from docx import Document
    except Exception as exc:
        return ExtractionResult("", "skipped", warnings=[f"python-docx unavailable: {exc}"])

    try:
        document = Document(str(record.path))
        lines: list[str] = []
        lines.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        for table_index, table in enumerate(document.tables, start=1):
            lines.append(f"TABLE {table_index}")
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return ExtractionResult(
            truncate_text(normalize_text("\n".join(lines))),
            "ok",
            metadata={"paragraphs": len(document.paragraphs), "tables": len(document.tables), "extractor": "docx"},
            warnings=[],
        )
    except Exception as exc:
        return ExtractionResult("", "error", warnings=[f"Cannot read DOCX: {exc}"])


def extract_rtf(record: FileRecord) -> ExtractionResult:
    try:
        from striprtf.striprtf import rtf_to_text
    except Exception as exc:
        return ExtractionResult("", "skipped", warnings=[f"striprtf unavailable: {exc}"])

    try:
        raw = record.path.read_text(encoding="utf-8", errors="ignore")
        return ExtractionResult(
            truncate_text(normalize_text(rtf_to_text(raw))),
            "ok",
            metadata={"extractor": "rtf"},
            warnings=[],
        )
    except Exception as exc:
        return ExtractionResult("", "error", warnings=[f"Cannot read RTF: {exc}"])


def extract_excel(record: FileRecord, max_rows: int) -> ExtractionResult:
    try:
        import pandas as pd
    except Exception as exc:
        return ExtractionResult("", "skipped", warnings=[f"pandas/openpyxl/xlrd unavailable for Excel: {exc}"])

    try:
        sheets = pd.read_excel(record.path, sheet_name=None, header=None, nrows=max_rows)
        lines: list[str] = []
        sheet_names: list[str] = []
        columns: set[str] = set()
        for sheet_name, df in sheets.items():
            sheet_names.append(str(sheet_name))
            lines.append(f"SHEET: {sheet_name}")
            for _, row in df.iterrows():
                values = ["" if value is None else str(value) for value in row.tolist()]
                for value in values:
                    if value and len(value) < 80:
                        columns.add(value)
                line = " | ".join(value for value in values if value and value.lower() != "nan")
                if line:
                    lines.append(line)
        return ExtractionResult(
            truncate_text(normalize_text("\n".join(lines))),
            "ok",
            metadata={"sheets": sheet_names, "columns_or_cells": sorted(list(columns))[:200], "extractor": "excel"},
            warnings=[],
        )
    except Exception as exc:
        return ExtractionResult("", "error", warnings=[f"Cannot read Excel: {exc}"])


def extract_doc(record: FileRecord) -> ExtractionResult:
    warnings: list[str] = []
    try:
        import textract
    except Exception as exc:
        warnings.append(f"textract unavailable: {exc.__class__.__name__}")
    else:
        try:
            data = textract.process(str(record.path))
            text = data.decode("utf-8", errors="ignore")
            return ExtractionResult(
                truncate_text(normalize_text(text)),
                "ok",
                metadata={"extractor": "textract-doc"},
                warnings=[],
            )
        except Exception as exc:
            warnings.append(f"textract failed: {exc}")

    try:
        completed = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(record.path)],
            check=False,
            capture_output=True,
            timeout=120,
        )
    except FileNotFoundError:
        warnings.append("textutil unavailable")
    except Exception as exc:
        warnings.append(f"textutil failed: {exc}")
    else:
        text = decode_textutil_output(completed.stdout)
        if completed.returncode == 0 and text.strip():
            return ExtractionResult(
                truncate_text(normalize_text(text)),
                "ok",
                metadata={"extractor": "textutil-doc"},
                warnings=warnings,
            )
        stderr = completed.stderr.decode("utf-8", errors="replace").strip()
        warnings.append(f"textutil returned no text: {stderr or f'exit_code={completed.returncode}'}")

    return ExtractionResult("", "skipped", warnings=[f"DOC extraction unavailable: {'; '.join(warnings)}"])


def decode_textutil_output(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    repaired = repair_cp1251_mojibake(text)
    if count_cyrillic_letters(repaired) > count_cyrillic_letters(text):
        return repaired
    return text


def repair_cp1251_mojibake(text: str) -> str:
    chunks: list[str] = []
    byte_buffer = bytearray()
    for char in text:
        codepoint = ord(char)
        if codepoint <= 255:
            byte_buffer.append(codepoint)
            continue
        if byte_buffer:
            chunks.append(byte_buffer.decode("cp1251", errors="replace"))
            byte_buffer.clear()
        chunks.append(char)
    if byte_buffer:
        chunks.append(byte_buffer.decode("cp1251", errors="replace"))
    return "".join(chunks)


def count_cyrillic_letters(text: str) -> int:
    return sum(1 for char in text if "А" <= char <= "я" or char in "Ёё")
